import io
import zipfile
from typing import TypedDict

from docx import Document


# DOCX has no reliable page count without rendering.
# Approximate at ~3500 chars/page (matches typical resume density).
_CHARS_PER_PAGE = 3500


class ParsedDocument(TypedDict):
    raw_text: str
    images: list[bytes]
    page_count: int
    author: str


def extract_docx(data: bytes) -> ParsedDocument:
    """Extract text and embedded images from a DOCX byte stream."""
    text, author = _extract_text_and_author(data)
    images = _extract_images(data)
    page_count = max(1, (len(text) + _CHARS_PER_PAGE - 1) // _CHARS_PER_PAGE)
    return {"raw_text": text, "images": images, "page_count": page_count, "author": author}


def _extract_text_and_author(data: bytes) -> tuple[str, str]:
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            parts.append(para.text)
    author = (doc.core_properties.author or "").strip()
    return "\n".join(parts), author


def _extract_images(data: bytes) -> list[bytes]:
    images: list[bytes] = []
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        for name in zf.namelist():
            if name.startswith("word/media/"):
                images.append(zf.read(name))
    return images
