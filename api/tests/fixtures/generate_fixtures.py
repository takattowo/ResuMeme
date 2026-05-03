"""Generates synthetic PDF and DOCX fixtures for parser tests.
Run once to regenerate; outputs are committed to the repo.
"""
import io
import os

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from PIL import Image
from docx import Document
from docx.shared import Inches

HERE = os.path.dirname(os.path.abspath(__file__))

SAMPLE_TEXT = {
    "header": "Jane Doe",
    "title": "Senior Software Engineer",
    "summary_heading": "Summary",
    "summary": "Experienced engineer with 8 years of building synergy.",
    "experience_heading": "Experience",
    "experience": [
        "Acme Corp - Staff Engineer (2020-2024)",
        "Built scalable systems leveraging cutting-edge stakeholder alignment.",
        "Globex - Senior Engineer (2016-2020)",
        "Spearheaded paradigm shifts in vertical-agnostic frameworks.",
    ],
    "skills_heading": "Skills",
    "skills": "Python, JavaScript, TypeScript, Go, Rust, Kubernetes, Synergy",
    "education_heading": "Education",
    "education": "BSc Computer Science, State University, 2016",
}


def _make_red_square(size: int = 80) -> bytes:
    img = Image.new("RGB", (size, size), color=(220, 40, 40))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def write_pdf(path: str, with_image: bool = False, with_sections: bool = True) -> None:
    c = canvas.Canvas(path, pagesize=letter)
    width, height = letter
    y = height - 60

    c.setFont("Helvetica-Bold", 18)
    c.drawString(60, y, SAMPLE_TEXT["header"])
    y -= 24
    c.setFont("Helvetica", 12)
    c.drawString(60, y, SAMPLE_TEXT["title"])
    y -= 36

    if with_image:
        img_bytes = _make_red_square()
        img = ImageReader(io.BytesIO(img_bytes))
        c.drawImage(img, 450, height - 140, width=80, height=80)

    if with_sections:
        for heading_key, body_key in [
            ("summary_heading", "summary"),
            ("experience_heading", "experience"),
            ("skills_heading", "skills"),
            ("education_heading", "education"),
        ]:
            c.setFont("Helvetica-Bold", 13)
            c.drawString(60, y, SAMPLE_TEXT[heading_key])
            y -= 18
            c.setFont("Helvetica", 11)
            body = SAMPLE_TEXT[body_key]
            if isinstance(body, list):
                for line in body:
                    c.drawString(60, y, line)
                    y -= 14
            else:
                c.drawString(60, y, body)
                y -= 14
            y -= 10
    else:
        c.setFont("Helvetica", 11)
        prose = (
            "I am a passionate engineer who loves to leverage synergy "
            "across diverse stakeholders to drive transformational outcomes."
        )
        c.drawString(60, y, prose)

    c.showPage()
    c.save()


def write_docx(path: str) -> None:
    doc = Document()
    doc.add_heading(SAMPLE_TEXT["header"], level=0)
    doc.add_paragraph(SAMPLE_TEXT["title"])

    for heading_key, body_key in [
        ("summary_heading", "summary"),
        ("experience_heading", "experience"),
        ("skills_heading", "skills"),
        ("education_heading", "education"),
    ]:
        doc.add_heading(SAMPLE_TEXT[heading_key], level=1)
        body = SAMPLE_TEXT[body_key]
        if isinstance(body, list):
            for line in body:
                doc.add_paragraph(line)
        else:
            doc.add_paragraph(body)

    img_path = os.path.join(HERE, "_temp_red.png")
    with open(img_path, "wb") as f:
        f.write(_make_red_square())
    doc.add_picture(img_path, width=Inches(1.0))
    os.remove(img_path)

    doc.save(path)


def main() -> None:
    write_pdf(os.path.join(HERE, "sample_cv.pdf"))
    write_pdf(os.path.join(HERE, "sample_cv_with_image.pdf"), with_image=True)
    write_pdf(
        os.path.join(HERE, "sample_cv_no_sections.pdf"),
        with_sections=False,
    )
    write_docx(os.path.join(HERE, "sample_cv.docx"))
    print("Fixtures generated.")


if __name__ == "__main__":
    main()
